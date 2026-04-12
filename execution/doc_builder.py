from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import zipfile
import shutil
from datetime import datetime
from typing import Any, Dict, List
from logger_config import setup_logger
import time

logger = setup_logger("DocBuilder")
CITATION_GROUP_PATTERN = re.compile(
    r'\[((?:\s*(?:Original(?:\s+Report)?|\d+)\s*)(?:,\s*(?:Original(?:\s+Report)?|\d+)\s*)*)\]',
    re.IGNORECASE,
)
VISUAL_MARKER_PATTERN = re.compile(
    r'\[(?:Figure|Asset|איור|תמונה|Figura|Image)(?:\s+ID)?\s*:?\s*([A-Za-z0-9][A-Za-z0-9_-]{3,63})[:\s]*.*?\]',
    re.IGNORECASE,
)
RAW_VISUAL_TOKEN_PATTERN = re.compile(r'\[\s*visual\s*:\s*([A-Za-z0-9][A-Za-z0-9_-]{3,63})[^\]]*\]', re.IGNORECASE)
_AMBIGUOUS_VISUAL = object()


def _reference_identity(ref: Dict) -> tuple:
    """Builds a stable identity for deduplicating references across chapters."""
    title = str(ref.get('title', '')).strip().lower()
    url = str(ref.get('url', '')).strip().lower()
    category = str(ref.get('category', '')).strip().lower()
    if url:
        return (title, url, category)
    return (title, category)


def _normalize_citation_token(value: Any) -> str:
    token = str(value or '').strip()
    if not token:
        return ''
    if token.isdigit():
        return str(int(token))
    if token.lower() in {'original', 'original report'}:
        return 'Original'
    return token


def _ordered_local_references(references: List[Dict]) -> List[Dict]:
    """Returns chapter-local references ordered by their declared index when present."""
    decorated = []
    for position, ref in enumerate(references or []):
        raw_index = _normalize_citation_token(ref.get('index'))
        if raw_index == 'Original':
            sort_index = float('inf')
        else:
            try:
                sort_index = int(raw_index)
            except (TypeError, ValueError):
                sort_index = position + 1
        decorated.append((sort_index, position, ref))
    decorated.sort(key=lambda item: (item[0], item[1]))
    return [ref for _, _, ref in decorated]


def _remap_inline_citations(text: str, citation_map: Dict[str, int]) -> str:
    """Remaps inline citations like [2] or [Original, 2] to global bibliography indices."""
    if not text or not citation_map:
        return text

    def replace(match):
        raw_tokens = [token.strip() for token in match.group(1).split(',')]
        remapped_tokens = []
        seen = set()

        for raw_token in raw_tokens:
            normalized = _normalize_citation_token(raw_token)
            if not normalized:
                continue
            mapped = citation_map.get(normalized)
            if mapped is None:
                logger.warning("Dropping unresolved citation token during export normalization: %s", normalized)
                continue
            mapped_token = str(mapped)
            if mapped_token in seen:
                continue
            seen.add(mapped_token)
            remapped_tokens.append(mapped_token)

        if not remapped_tokens:
            return ''

        return f"[{', '.join(remapped_tokens)}]"

    return CITATION_GROUP_PATTERN.sub(replace, text)


def _remap_string_list(items: List[str], citation_map: Dict[str, int]) -> List[str]:
    remapped = []
    for item in items or []:
        text = _remap_inline_citations(item, citation_map).strip()
        if text:
            remapped.append(text)
    return remapped


def _remap_export_text_fields(chapter_copy: Dict[str, Any], citation_map: Dict[str, int]) -> None:
    for field in ("draft_text", "executive_takeaway"):
        chapter_copy[field] = _remap_inline_citations(chapter_copy.get(field, ''), citation_map)

    for field in ("retained_claims", "updated_claims", "new_claims", "open_questions"):
        chapter_copy[field] = _remap_string_list(chapter_copy.get(field, []), citation_map)


def normalize_report_citations(chapters: List[Dict]) -> tuple[List[Dict], List[Dict]]:
    """
    Produces a citation-normalized copy of the chapters and an ordered global bibliography.
    Chapter-local references are remapped to global indices and inline citations in draft_text
    are rewritten to match the final bibliography ordering.
    """
    normalized_chapters = []
    ordered_refs = []
    global_ref_map = {}

    for chapter in chapters:
        chapter_copy = chapter.copy()
        chapter_refs = _ordered_local_references(chapter.get('references', []))
        local_to_global: Dict[str, int] = {}
        normalized_refs = []

        for fallback_index, ref in enumerate(chapter_refs, start=1):
            local_index = _normalize_citation_token(ref.get('index', fallback_index))
            if not local_index:
                local_index = str(fallback_index)

            ref_copy = dict(ref)
            ref_key = _reference_identity(ref_copy)
            if ref_key not in global_ref_map:
                global_index = len(ordered_refs) + 1
                global_ref_map[ref_key] = global_index
                ref_copy['index'] = global_index
                ordered_refs.append(ref_copy)
            else:
                global_index = global_ref_map[ref_key]

            local_to_global[local_index] = global_index
            local_to_global.setdefault(str(fallback_index), global_index)
            ref_copy['index'] = global_index
            normalized_refs.append(ref_copy)

        _remap_export_text_fields(chapter_copy, local_to_global)
        chapter_copy['references'] = normalized_refs
        normalized_chapters.append(chapter_copy)

    return normalized_chapters, ordered_refs


def _normalize_marker_token(value: Any) -> str:
    token = str(value or "").strip().lower()
    if not token:
        return ''
    return re.sub(r'[^a-z0-9_-]', '', token)


def _register_short_visual(mapping: Dict[str, Any], key: str, visual: Dict[str, Any]) -> None:
    if not key:
        return
    existing = mapping.get(key)
    if existing is None:
        mapping[key] = visual
    elif existing is not visual:
        mapping[key] = _AMBIGUOUS_VISUAL


def _build_visual_lookup(visuals: List[Dict] | None) -> Dict[str, Dict[str, Any]]:
    exact_map: Dict[str, Dict[str, Any]] = {}
    short_original_map: Dict[str, Any] = {}
    short_generated_map: Dict[str, Any] = {}

    for visual in visuals or []:
        marker_id = _normalize_marker_token(visual.get('marker_id'))
        visual_id = _normalize_marker_token(visual.get('id'))
        original_asset_id = _normalize_marker_token(visual.get('original_asset_id'))

        for token in (marker_id, visual_id, original_asset_id):
            if token:
                exact_map[token] = visual

        if original_asset_id:
            _register_short_visual(short_original_map, original_asset_id[:8], visual)
        if visual_id:
            _register_short_visual(short_generated_map, visual_id[:8], visual)
        if marker_id:
            _register_short_visual(short_generated_map, marker_id[:8], visual)

    return {
        'exact': exact_map,
        'short_original': short_original_map,
        'short_generated': short_generated_map,
    }


def _resolve_visual(lookup: Dict[str, Dict[str, Any]], marker_token: str) -> Dict[str, Any] | None:
    normalized = _normalize_marker_token(marker_token)
    if not normalized:
        return None

    exact = lookup['exact'].get(normalized)
    if exact is not None:
        return exact

    short_token = normalized[:8]
    for bucket_name in ('short_original', 'short_generated'):
        candidate = lookup[bucket_name].get(short_token)
        if candidate is _AMBIGUOUS_VISUAL:
            return None
        if candidate is not None:
            return candidate
    return None


def _visual_export_id(visual: Dict[str, Any], fallback: str) -> str:
    for key in ('marker_id', 'original_asset_id', 'id'):
        token = _normalize_marker_token(visual.get(key))
        if token:
            return token
    return _normalize_marker_token(fallback) or 'visual'


def _unique_ordered(items: List[str], limit: int = 10) -> List[str]:
    seen = set()
    results = []
    for item in items:
        text = str(item).strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        results.append(text)
        if len(results) >= limit:
            break
    return results


def _contains_hebrew(text: str) -> bool:
    return bool(re.search(r'[\u0590-\u05FF]', text or ''))


def _resolve_output_language(report_metadata: Dict | None, report_title: str, chapters: List[Dict]) -> str:
    report_metadata = report_metadata or {}
    requested = str(report_metadata.get("output_language", "Match Source Language")).strip()
    if requested in {"English", "Hebrew"}:
        return requested

    sample_parts = [report_title]
    for chapter in chapters[:3]:
        sample_parts.append(chapter.get("title", ""))
        sample_parts.append(chapter.get("draft_text", "")[:200])
    sample_text = " ".join(part for part in sample_parts if part)
    return "Hebrew" if _contains_hebrew(sample_text) else "English"


def _language_pack(output_language: str) -> Dict[str, str]:
    if output_language == "Hebrew":
        return {
            "cover_subtitle": "מהדורה מעודכנת",
            "date_label": "תאריך",
            "table_of_contents": "תוכן עניינים",
            "executive_summary": "תקציר מנהלים",
            "methodology": "מתודולוגיה",
            "key_developments": "התפתחויות מרכזיות מאז המהדורה המקורית",
            "updated_chapters": "פרקים מעודכנים",
            "bibliography": "ביבליוגרפיה",
            "metadata_summary": "תאריך המהדורה המקורית: {original_report_date} | מעודכן עד: {update_end_date} | קהל יעד: {target_audience}",
            "executive_intro": "מהדורה מעודכנת זו של {report_title} מרעננת את הדוח המקורי מתאריך {original_report_date} עם ממצאים והתפתחויות עד {update_end_date}. הדוח נערך עבור {target_audience} ושומר על קווי החשיבה החזקים של הדוח המקורי תוך עדכון בסיס הראיות.",
            "methodology_intro": "דוח זה נבנה באמצעות שילוב בין טקסט הדוח המקורי, ממצאי מחקר מעודכנים, חומרי ייחוס שהועלו, חזותיים מעודכנים ותהליך סופי של נרמול ציטוטים לפני הייצוא.",
            "methodology_original": "הדוח המקורי שימש כבסיס למסגור היסטורי ולטענות שעדיין נותרו תקפות.",
            "methodology_window": "המחקר תועדף כך שיכסה את חלון העדכון שנבחר עד {update_end_date}.",
            "methodology_citations": "הציטוטים המוטמעים נורמלו גלובלית כדי שמספור הביבליוגרפיה הסופי יישאר עקבי לאורך כל הדוח.",
            "methodology_categories": "קטגוריות המקורות במהדורה זו: {categories}.",
            "empty_executive": "לא היו זמינות תובנות מנהלים ברמת הפרק, ולכן התקציר נבנה ממבנה הפרקים.",
            "empty_developments": "לא היו זמינות נקודות מפורשות על התפתחויות מרכזיות מתוך פלט הפרקים.",
        }

    return {
        "cover_subtitle": "Updated Edition",
        "date_label": "Date",
        "table_of_contents": "Table of Contents",
        "executive_summary": "Executive Summary",
        "methodology": "Methodology",
        "key_developments": "Key Developments Since the Original Edition",
        "updated_chapters": "Updated Chapters",
        "bibliography": "Bibliography",
        "metadata_summary": "Original edition date: {original_report_date} | Updated through: {update_end_date} | Audience: {target_audience}",
        "executive_intro": "This updated edition of {report_title} refreshes the original report dated {original_report_date} with evidence and developments through {update_end_date}. The report is organized for {target_audience} and preserves the original report's strongest through-lines while updating the evidence base.",
        "methodology_intro": "This report was generated by combining the original report text with refreshed research findings, uploaded reference materials, updated figures, and a final citation-normalization pass before export.",
        "methodology_original": "The original report was treated as the baseline source for historical framing and still-valid claims.",
        "methodology_window": "Research was prioritized to cover the selected update window through {update_end_date}.",
        "methodology_citations": "Inline citations were normalized globally so the final bibliography numbering stays consistent across chapters.",
        "methodology_categories": "Source categories used in this edition: {categories}.",
        "empty_executive": "Chapter-level executive takeaways were not available, so this summary was generated from the chapter structure.",
        "empty_developments": "No explicit key-development bullets were available from the chapter outputs.",
    }


def synthesize_report_sections(chapters: List[Dict], report_title: str, report_metadata: Dict | None = None) -> Dict[str, List[str] | str]:
    report_metadata = report_metadata or {}
    original_report_date = report_metadata.get("original_report_date", "Unknown")
    update_end_date = report_metadata.get("update_end_date", datetime.now().strftime("%Y-%m-%d"))
    target_audience = report_metadata.get("target_audience", "General professional audience")
    output_language = _resolve_output_language(report_metadata, report_title, chapters)
    language_pack = _language_pack(output_language)

    executive_points = []
    key_developments = []
    source_categories = []

    for chapter in chapters:
        chapter_title = chapter.get("title", "Untitled Chapter")
        takeaway = chapter.get("executive_takeaway")
        if takeaway:
            executive_points.append(f"{chapter_title}: {takeaway}")

        for claim in chapter.get("new_claims", []):
            key_developments.append(f"{chapter_title}: {claim}")
        for claim in chapter.get("updated_claims", []):
            key_developments.append(f"{chapter_title}: {claim}")

        for ref in chapter.get("references", []):
            category = ref.get("category")
            if category:
                source_categories.append(category)

    executive_points = _unique_ordered(executive_points, limit=6)
    key_developments = _unique_ordered(key_developments, limit=10)
    source_categories = _unique_ordered(source_categories, limit=10)

    metadata_summary = language_pack["metadata_summary"].format(
        original_report_date=original_report_date,
        update_end_date=update_end_date,
        target_audience=target_audience,
    )
    executive_intro = language_pack["executive_intro"].format(
        report_title=report_title,
        original_report_date=original_report_date,
        update_end_date=update_end_date,
        target_audience=target_audience,
    )
    methodology_intro = language_pack["methodology_intro"]
    methodology_points = [
        language_pack["methodology_original"],
        language_pack["methodology_window"].format(update_end_date=update_end_date),
        language_pack["methodology_citations"],
    ]
    if source_categories:
        methodology_points.append(language_pack["methodology_categories"].format(categories=', '.join(source_categories)))

    if not executive_points:
        executive_points = [language_pack["empty_executive"]]
    if not key_developments:
        key_developments = [language_pack["empty_developments"]]

    return {
        "output_language": output_language,
        "titles": {
            "table_of_contents": language_pack["table_of_contents"],
            "executive_summary": language_pack["executive_summary"],
            "methodology": language_pack["methodology"],
            "key_developments": language_pack["key_developments"],
            "updated_chapters": language_pack["updated_chapters"],
            "bibliography": language_pack["bibliography"],
        },
        "cover_subtitle": language_pack["cover_subtitle"],
        "date_label": language_pack["date_label"],
        "metadata_summary": metadata_summary,
        "executive_intro": executive_intro,
        "executive_points": executive_points,
        "methodology_intro": methodology_intro,
        "methodology_points": methodology_points,
        "key_developments": key_developments,
    }

class ReportBuilder:
    def __init__(self, title: str):
        logger.info(f"Initializing ReportBuilder for report: '{title}'")
        self.doc = Document()
        self.title = title
        self._chapter_count = 0  # Track chapters to add section breaks
        self._setup_styles()
        self._add_page_numbers()

    def _setup_styles(self):
        # Could customize default styles here if needed
        pass

    def _add_page_numbers(self):
        """Adds a right-aligned page number to the footer of the default section."""
        try:
            section = self.doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.clear()  # Remove any existing content

            run = para.add_run()
            # Add the PAGE field using raw XML
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld)

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._r.append(instrText)

            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_end)

            logger.debug("Page numbers added to footer.")
        except Exception as e:
            logger.warning(f"Could not add page numbers: {e}")

    def add_cover_page(self, subtitle: str, metadata_summary: str = "", date_label: str = "Date"):
        logger.info(f"Adding cover page with subtitle: '{subtitle}'")
        title_para = self.doc.add_heading(self.title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_para = self.doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if metadata_summary:
            metadata_para = self.doc.add_paragraph(metadata_summary)
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_rtl_if_hebrew(metadata_para, metadata_summary)
        
        date_para = self.doc.add_paragraph(f"{date_label}: {datetime.now().strftime('%Y-%m-%d')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()

    def add_table_of_contents(self, title: str = "Table of Contents"):
        """Inserts a Word Table of Contents field that auto-updates when the doc is opened."""
        logger.info("Adding Table of Contents field.")
        toc_heading = self.doc.add_heading(title, level=1)
        self._set_rtl_if_hebrew(toc_heading, self.title)  # Match document direction

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        # BEGIN field
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        # Field instruction
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r' TOC \o "1-3" \h \z \u '
        run._r.append(instrText)
        # SEPARATE field
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar_sep)
        # Placeholder text shown before user updates TOC in Word
        run2 = paragraph.add_run("Right-click and select 'Update Field' to generate table of contents.")
        run2.font.color.rgb = None  # Default color
        # END field
        run3 = paragraph.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar_end)

        # Add a setting to auto-update fields when the document is opened
        try:
            settings = self.doc.settings.element
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            settings.append(update_fields)
        except Exception as e:
            logger.warning(f"Could not set auto-update fields: {e}")

        self.doc.add_page_break()
        logger.info("Table of Contents field added.")

    def add_chapter(self, title: str, text: str, visuals: List[Dict] = None):
        logger.info(f"Adding chapter: '{title}' (text length: {len(text)})")
        # Add page break before each chapter (except the very first after the cover)
        if self._chapter_count > 0:
            self.doc.add_page_break()
        self._chapter_count += 1

        h = self.doc.add_heading(title, level=1)
        self._set_rtl_if_hebrew(h, title)
        
        visual_lookup = _build_visual_lookup(visuals)
        if visuals:
            for v in visuals:
                logger.debug(
                    "Registered visual for export: marker=%s id=%s original=%s type=%s path=%s",
                    v.get('marker_id'),
                    v.get('id'),
                    v.get('original_asset_id'),
                    v.get('type'),
                    v.get('path'),
                )

        # Strip placeholder markers (e.g. [Figure XXXXXXXX: Caption]) that the LLM
        # emits by copying the example literally — these can never be matched to a visual.
        placeholder_regex = r'\[(?:Figure|Asset|איור|תמונה|Figura|Image)(?:\s+ID)?\s*:?\s*X{4,}[:\s]*.*?\]'
        text = re.sub(placeholder_regex, '', text, flags=re.IGNORECASE)
        text = RAW_VISUAL_TOKEN_PATTERN.sub('', text)

        # Split text into lines to handle headers and lists properly
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle visual markers first, including legacy short IDs and canonical marker IDs.
            marker_pattern = VISUAL_MARKER_PATTERN
            if marker_pattern.search(line):
                # If we have a marker in the line, we process it specifically
                # We use a non-capturing group for the inner elements and capture the WHOLE tag to split cleanly
                split_pattern = r'(\[(?:Figure|Asset|איור|תמונה|Figura|Image)(?:\s+ID)?\s*:?\s*[A-Za-z0-9][A-Za-z0-9_-]{3,63}[:\s]*.*?\])'
                parts = re.split(split_pattern, line, flags=re.IGNORECASE)
                p = self.doc.add_paragraph()
                for part in parts:
                    if not part: continue
                    match = marker_pattern.match(part)
                    if match:
                        marker_token = match.group(1)
                        asset_id_short = _normalize_marker_token(marker_token)[:8]
                        v = _resolve_visual(visual_lookup, marker_token)
                        img_path = v.get('path') if v else None
                        cap_text = f"Figure: {v.get('title', v.get('short_caption', 'Visual'))}" if v else f"Figure {asset_id_short}"

                        if v is not None and not (img_path and os.path.exists(img_path)):
                            import glob
                            search_token = _visual_export_id(v, marker_token)
                            possible = glob.glob(f".tmp/**/*{search_token}*.*", recursive=True)
                            for p_path in possible:
                                if p_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                                    img_path = p_path
                                    break

                        if img_path and os.path.exists(img_path):
                            try:
                                img_p = self.doc.add_paragraph()
                                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                img_p.add_run().add_picture(img_path, width=Inches(5))
                                cap = self.doc.add_paragraph(cap_text)
                                cap.style = 'Caption'
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                self._set_rtl_if_hebrew(cap, cap_text)
                            except Exception as e:
                                logger.error(f"Failed to add picture {img_path} to document: {e}")
                                pass
                        else:
                            logger.warning(f"Visual marker found for ID '{asset_id_short}' but no matching visual found.")
                            continue
                    else:
                        self._add_markdown_runs(p, part)
                self._set_rtl_if_hebrew(p, line)
                continue

            # Handle Headers
            header_match = re.match(r'^(#{1,6})\s*(.*)', line)
            if header_match:
                level = len(header_match.group(1))
                h_text = header_match.group(2)
                h = self.doc.add_heading(h_text, level=min(level + 1, 9))
                self._set_rtl_if_hebrew(h, h_text)
                continue

            # Handle List Items
            list_match = re.match(r'^[\*\-\+]\s*(.*)', line)
            if list_match:
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_markdown_runs(p, list_match.group(1))
                self._set_rtl_if_hebrew(p, line)
                continue
                
            num_list_match = re.match(r'^\d+\.\s*(.*)', line)
            if num_list_match:
                p = self.doc.add_paragraph(style='List Number')
                self._add_markdown_runs(p, num_list_match.group(1))
                self._set_rtl_if_hebrew(p, line)
                continue

            # Normal Paragraph
            p = self.doc.add_paragraph()
            self._add_markdown_runs(p, line)
            self._set_rtl_if_hebrew(p, line)

    def _add_markdown_runs(self, paragraph, text):
        """
        Parses bold and italic markdown and adds them as runs to a paragraph.
        """
        # Pattern for bold (***bolditalic***, **bold**, *italic*)
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__| _.*?_)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part: continue
            
            run = paragraph.add_run()
            if part.startswith('***') and part.endswith('***'):
                run.text = part[3:-3]
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run.text = part[2:-2]
                run.bold = True
            elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                run.text = part[1:-1]
                run.italic = True
            else:
                run.text = part

    def _set_rtl_if_hebrew(self, paragraph, text):
        # Check for Hebrew characters — apply RTL silently (no per-paragraph log)
        if re.search(r'[\u0590-\u05FF]', text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.bidi = True
            for run in paragraph.runs:
                run.font.rtl = True

    def add_bibliography(self, references: List[Dict], title: str = "Bibliography"):
        """Adds a numbered bibliography. References are rendered as [index] Title - URL,
        matching the numbered inline citations [1], [2], etc. in the report text."""
        logger.info(f"Adding bibliography with {len(references)} references.")
        self.doc.add_page_break()
        bib_heading = self.doc.add_heading(title, level=1)
        self._set_rtl_if_hebrew(bib_heading, title)

        for ref in references:
            idx = ref.get('index', references.index(ref) + 1)
            title = ref.get('title', 'Unknown Source')
            url = ref.get('url', '')
            category = ref.get('category', '')

            p = self.doc.add_paragraph()
            # Index number in bold
            idx_run = p.add_run(f"[{idx}] ")
            idx_run.bold = True
            # Title
            title_run = p.add_run(title)
            title_run.bold = True
            # URL if available
            if url and url.startswith('http'):
                p.add_run(f" — {url}")
            # Category tag
            if category:
                cat_run = p.add_run(f"  [{category}]")
                cat_run.italic = True
            self._set_rtl_if_hebrew(p, title)

    def add_front_matter_section(self, title: str, intro_text: str = "", bullets: List[str] | None = None):
        logger.info(f"Adding front matter section: '{title}'")
        heading = self.doc.add_heading(title, level=1)
        self._set_rtl_if_hebrew(heading, title)

        if intro_text:
            paragraph = self.doc.add_paragraph(intro_text)
            self._set_rtl_if_hebrew(paragraph, intro_text)

        for bullet in bullets or []:
            paragraph = self.doc.add_paragraph(style='List Bullet')
            self._add_markdown_runs(paragraph, bullet)
            self._set_rtl_if_hebrew(paragraph, bullet)

    def save(self, output_path: str):
        save_start = time.time()
        logger.info(f"Saving DOCX to: {output_path}")
        self.doc.save(output_path)
        logger.info(f"DOCX saved in {time.time() - save_start:.2f}s.")
        return output_path

def build_final_report(chapters: List[Dict], output_path: str = "Modernized_Report.docx", title: str = None, report_metadata: Dict | None = None) -> str:
    report_title = title or "Modernized Industry Report"
    logger.info(f"Starting build_final_report for {len(chapters)} chapters. Title: '{report_title}'")
    normalized_chapters, ordered_refs = normalize_report_citations(chapters)
    sections = synthesize_report_sections(normalized_chapters, report_title, report_metadata)
    include_front_matter = bool((report_metadata or {}).get("include_synthetic_front_matter"))
    builder = ReportBuilder(report_title)
    builder.add_cover_page(sections["cover_subtitle"], metadata_summary=sections["metadata_summary"], date_label=sections["date_label"])
    builder.add_table_of_contents(sections["titles"]["table_of_contents"])
    if include_front_matter:
        builder.add_front_matter_section(sections["titles"]["executive_summary"], sections["executive_intro"], sections["executive_points"])
        builder.add_front_matter_section(sections["titles"]["methodology"], sections["methodology_intro"], sections["methodology_points"])
        builder.add_front_matter_section(sections["titles"]["key_developments"], bullets=sections["key_developments"])

    for ch in normalized_chapters:
        builder.add_chapter(ch['title'], ch['draft_text'], ch.get('approved_visuals', []))

    builder.add_bibliography(ordered_refs, title=sections["titles"]["bibliography"])
    
    return builder.save(output_path)

def build_markdown_report(chapters: List[Dict], output_path: str = "Modernized_Report.md", title: str = None, report_metadata: Dict | None = None) -> str:
    report_title = title or "Modernized Industry Report"
    logger.info(f"Starting build_markdown_report for {len(chapters)} chapters. Title: '{report_title}'")
    normalized_chapters, ordered_refs = normalize_report_citations(chapters)
    sections = synthesize_report_sections(normalized_chapters, report_title, report_metadata)
    include_front_matter = bool((report_metadata or {}).get("include_synthetic_front_matter"))

    md_content = []
    
    # Create export directory for markdown and images
    base_name = os.path.splitext(output_path)[0]
    export_dir = f"{base_name}_export"
    images_dir = os.path.join(export_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    
    output_md_path = os.path.join(export_dir, os.path.basename(output_path))
    
    # Cover & TOC omitted in simple markdown, just title
    md_content.append(f"# {report_title}\n")
    md_content.append(f"_{sections['cover_subtitle']} | {sections['date_label']}: {datetime.now().strftime('%Y-%m-%d')}_\n\n")
    md_content.append(f"**{sections['metadata_summary']}**\n")
    if include_front_matter:
        md_content.append(f"\n## {sections['titles']['executive_summary']}\n")
        md_content.append(f"{sections['executive_intro']}\n")
        for point in sections["executive_points"]:
            md_content.append(f"- {point}")
        md_content.append(f"\n## {sections['titles']['methodology']}\n")
        md_content.append(f"{sections['methodology_intro']}\n")
        for point in sections["methodology_points"]:
            md_content.append(f"- {point}")
        md_content.append(f"\n## {sections['titles']['key_developments']}\n")
        for point in sections["key_developments"]:
            md_content.append(f"- {point}")

    for ch in normalized_chapters:
        md_content.append(f"# {ch['title']}\n")
        
        visual_lookup = _build_visual_lookup(ch.get('approved_visuals', []))
        
        text = ch['draft_text']
        placeholder_regex = r'\[(?:Figure|Asset|איור|תמונה|Figura|Image)(?:\s+ID)?\s*:?\s*X{4,}[:\s]*.*?\]'
        text = re.sub(placeholder_regex, '', text, flags=re.IGNORECASE)
        text = RAW_VISUAL_TOKEN_PATTERN.sub('', text)
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                md_content.append("")
                continue
                
            marker_pattern = VISUAL_MARKER_PATTERN
            if marker_pattern.search(line):
                split_pattern = r'(\[(?:Figure|Asset|איור|תמונה|Figura|Image)(?:\s+ID)?\s*:?\s*[A-Za-z0-9][A-Za-z0-9_-]{3,63}[:\s]*.*?\])'
                parts = re.split(split_pattern, line, flags=re.IGNORECASE)
                out_line = ""
                for part in parts:
                    if not part: continue
                    match = marker_pattern.match(part)
                    if match:
                        marker_token = match.group(1)
                        asset_id_short = _normalize_marker_token(marker_token)[:8]
                        v = _resolve_visual(visual_lookup, marker_token)
                        img_path = v.get('path') if v else None
                        cap_text = f"Figure: {v.get('title', v.get('short_caption', 'Visual'))}" if v else f"Figure {asset_id_short}"

                        if v is not None and not (img_path and os.path.exists(img_path)):
                            import glob
                            search_token = _visual_export_id(v, marker_token)
                            possible = glob.glob(f".tmp/**/*{search_token}*.*", recursive=True)
                            for p_path in possible:
                                if p_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                                    img_path = p_path
                                    break

                        if img_path and os.path.exists(img_path):
                            try:
                                img_ext = os.path.splitext(img_path)[1]
                                if not img_ext:
                                    img_ext = ".png" # fallback
                                dest_img_name = f"{_visual_export_id(v, marker_token)}{img_ext}"
                                dest_img_path = os.path.join(images_dir, dest_img_name)
                                shutil.copy2(img_path, dest_img_path)
                                rel_path = f"images/{dest_img_name}"
                            except Exception as e:
                                logger.error(f"Error copying image {img_path}: {e}")
                                rel_path = img_path # fallback
                            out_line += f"\n\n![{cap_text}]({rel_path})\n*{cap_text}*\n\n"
                        else:
                            logger.warning(f"Skipping unmatched visual marker '{asset_id_short}' in markdown export.")
                            continue
                    else:
                        out_line += part
                md_content.append(out_line)
                continue
            
            md_content.append(line)
        
        md_content.append("\n---\n")
        
    if ordered_refs:
        md_content.append(f"# {sections['titles']['bibliography']}\n")
        for ref in ordered_refs:
            idx = ref.get('index', ordered_refs.index(ref) + 1)
            title = ref.get('title', 'Unknown Source')
            url = ref.get('url', '')
            category = ref.get('category', '')
            ref_line = f"**[{idx}] {title}**"
            if url and url.startswith('http'):
                ref_line += f" — {url}"
            if category:
                ref_line += f"  _({category})_"
            md_content.append(ref_line)
            
    try:
        with open(output_md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))
            
        zip_output_path = f"{base_name}.zip"
        with zipfile.ZipFile(zip_output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add markdown file
            zipf.write(output_md_path, arcname=os.path.basename(output_md_path))
            
            # Add images
            for root, _, files in os.walk(images_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = f"images/{file}"
                    zipf.write(file_path, arcname=arcname)
                    
        logger.info(f"Markdown report zipped and saved to: {zip_output_path}")
        return zip_output_path
    except Exception as e:
        logger.error(f"Failed to generate Markdown report: {e}")
        return output_path

if __name__ == "__main__":
    test_chapters = [
        {
            "title": "Introduction to AI", 
            "draft_text": "# Header 1\n## Header 2\n**Bold Text** and *Italic Text*.\n- Item 1\n- Item 2",
            "approved_visuals": [],
            "references": [{"title": "Future of AI", "url": "http://example.com", "category": "Academic"}]
        }
    ]
    path = build_final_report(test_chapters, "test_markdown.docx")
    print(f"Report saved to {path}")
