from __future__ import annotations

import os
import sys
from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parents[1]
EXECUTION_ROOT = REPO_ROOT / "execution"

for candidate in (str(REPO_ROOT), str(EXECUTION_ROOT)):
    if candidate not in sys.path:
        sys.path.insert(0, candidate)


@pytest.fixture(autouse=True)
def no_live_gemini(monkeypatch, request):
    is_live_api_test = request.node.get_closest_marker("live_api") is not None
    if is_live_api_test and os.getenv("RUN_LIVE_API_TESTS") == "1":
        return
    monkeypatch.delenv("GEMINI_API_KEY", raising=False)


@pytest.fixture
def report_metadata() -> dict[str, str]:
    return {
        "original_report_date": "2023-03-15",
        "update_start_date": "2024-01-01",
        "update_end_date": "2026-03-22",
        "target_audience": "Executives",
        "output_language": "English",
    }


@pytest.fixture
def dummy_image_path(tmp_path: Path) -> Path:
    path = tmp_path / "dummy.png"
    Image.new("RGB", (320, 240), color="navy").save(path)
    return path


@pytest.fixture
def sample_visual(dummy_image_path: Path) -> dict[str, str]:
    return {
        "id": "abcdef1234567890",
        "type": "image",
        "path": str(dummy_image_path),
        "title": "Dummy Visual",
        "short_caption": "Dummy Visual",
    }


@pytest.fixture
def sample_txt_path(tmp_path: Path) -> Path:
    path = tmp_path / "reference.txt"
    path.write_text(
        "Report date: March 15 2024.\n\n"
        "AI adoption accelerated sharply in 2024 as enterprise spending rose by 20 percent.\n\n"
        "A second paragraph discusses regulation and governance in 2025.\n\n"
        "A final paragraph covers product launches and deployment risks in 2026.\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path, dummy_image_path: Path) -> Path:
    path = tmp_path / "reference.docx"
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Report date: March 15 2024.")
    doc.add_paragraph("AI adoption accelerated in 2024 and spending rose by 20 percent.")
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(dummy_image_path))
    doc.save(path)
    return path


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    path = tmp_path / "reference.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Executive Summary", fontsize=30)
    page.insert_textbox(
        fitz.Rect(72, 120, 520, 320),
        (
            "Report date: March 15 2024.\n"
            "AI adoption accelerated in 2024 as enterprises increased spending by 20 percent.\n"
            "The section remains readable and searchable for internal reference tests."
        ),
        fontsize=11,
    )
    doc.save(path)
    doc.close()
    return path
